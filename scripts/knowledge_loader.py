# -*- coding: utf-8 -*-
"""
知识库读取模块
支持格式: .txt / .md / .pdf / .docx
按 token 上限截断，避免上下文超长。
"""
import os

def _read_text_file(path: str) -> str:
    for enc in ("utf-8", "gbk", "utf-8-sig"):
        try:
            with open(path, "r", encoding=enc) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
    # 最后兜底
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def _read_pdf(path: str) -> str:
    from pypdf import PdfReader
    reader = PdfReader(path)
    parts = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            parts.append(text)
    return "\n\n".join(parts)


def _read_docx(path: str) -> str:
    import docx
    doc = docx.Document(path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    # 表格内容
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            parts.append(" | ".join(cells))
    return "\n".join(parts)


# 图片 OCR（离线，数据不出本机）
_ocr_engine = None

def _get_ocr():
    global _ocr_engine
    if _ocr_engine is None:
        from rapidocr_onnxruntime import RapidOCR
        _ocr_engine = RapidOCR()
    return _ocr_engine


def _read_image(path: str) -> str:
    """用离线 OCR 把截图/图片中的文字识别出来。超大图自动降采样，避免内存爆炸与识别超时。"""
    from PIL import Image
    Image.MAX_IMAGE_PIXELS = 200_000_000  # 放宽到 2 亿像素，95M 像素的思维导图不再报警
    engine = _get_ocr()
    ocr_path = path
    tmp_path = None
    try:
        with Image.open(path) as im:
            w, h = im.size
            max_side = 4096
            if max(w, h) > max_side:
                scale = max_side / max(w, h)
                im = im.resize((max(1, int(w * scale)), max(1, int(h * scale))), Image.LANCZOS)
                tmp_path = os.path.splitext(path)[0] + ".__ocr_tmp.png"
                im.save(tmp_path)
                ocr_path = tmp_path
    except Exception:
        pass  # 非图片或解析失败，交给 OCR 引擎自行尝试
    try:
        result, _ = engine(ocr_path)
        if not result:
            return ""
        lines = [item[1] for item in result if item[1] and item[1].strip()]
        return "\n".join(lines)
    finally:
        if tmp_path and os.path.exists(tmp_path):
            try:
                os.remove(tmp_path)
            except Exception:
                pass  # 沙箱回收站不可用时忽略，临时文件残留无碍


def read_file(path: str) -> str:
    ext = os.path.splitext(path)[1].lower()
    if ext in (".txt", ".md", ".markdown"):
        return _read_text_file(path)
    if ext == ".pdf":
        return _read_pdf(path)
    if ext in (".docx", ".doc"):
        return _read_docx(path)
    if ext in (".png", ".jpg", ".jpeg", ".bmp", ".webp"):
        return _read_image(path)
    # 未知格式按文本尝试
    return _read_text_file(path)


def load_knowledge_dir(dir_path: str, max_chars: int = 20000) -> str:
    """读取目录下所有支持的文件，合并成一段文本（按大小截断）。"""
    if not dir_path or not os.path.isdir(dir_path):
        return ""

    texts = []
    for root, _, files in os.walk(dir_path):
        for fname in sorted(files):
            if fname.startswith(".") or fname.startswith("~$") or "__ocr_tmp" in fname:
                continue
            fpath = os.path.join(root, fname)
            try:
                text = read_file(fpath)
                if text.strip():
                    texts.append(f"【文件: {os.path.relpath(fpath, dir_path)}】\n{text}")
            except Exception as e:
                texts.append(f"【文件: {fname} 读取失败: {e}】")

    merged = "\n\n" + "\n\n".join(texts) if texts else ""
    if len(merged) > max_chars:
        merged = merged[:max_chars] + "\n\n……（知识库内容过长，已截断）"
    return merged
